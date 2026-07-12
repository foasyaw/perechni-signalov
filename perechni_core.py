# -*- coding: utf-8 -*-
"""
Ядро: извлечение сигналов из чертежей схем подключения (DWG/DXF)
и заполнение Word-перечней входных/выходных сигналов по шаблону.

Логика проверена на реальном проекте.
Ожидаемый стиль чертежа:
  - заголовки аналоговых модулей: TEXT "Модуль X.Y AI|AO|WI ..." ;
  - описания каналов: MTEXT на слое "Текст" с "поз. <ТЕГ>." ;
  - метки БИЗ: слой "mark2" (наличие -> Exia, иначе Exd);
  - дискретные каналы: якоря "X.Y-KLDI<n>" / "X.Y-KLDO<n>"
    на слоях "Реле DI (1-KL)" / "Реле DO (3-KL)".
Шаблон Word: как фирменные перечни (СОДЕРЖАНИЕ-таблица, аналоговая и
дискретная таблицы по 7 колонок, ЛИСТ РЕГИСТРАЦИИ, разрыв секции после
содержания, разрывы страниц между таблицами).
"""
import copy, os, re, shutil, subprocess, tempfile, statistics, collections

import ezdxf
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- DWG -> DXF
# Два движка: AutoCAD Core Console (accoreconsole.exe, идёт с AutoCAD)
# и ODA File Converter (бесплатный, opendesign.com). Находится любой из них.

DEFAULT_ODA = r"E:\Конвертер автокад\ODAFileConverter.exe"

def find_oda():
    if os.path.exists(DEFAULT_ODA):
        return DEFAULT_ODA
    for root in (r"C:\Program Files\ODA", r"C:\Program Files (x86)\ODA"):
        if os.path.isdir(root):
            for dirpath, _dirs, files in os.walk(root):
                if "ODAFileConverter.exe" in files:
                    return os.path.join(dirpath, "ODAFileConverter.exe")
    return ""

def find_accore():
    roots = [r"C:\Program Files\Autodesk"]
    # диски C..G, папки вида "AutoCAD 20xx"
    for drive in "CDEFG":
        roots.append(drive + ":\\")
    seen = []
    for root in roots:
        if not os.path.isdir(root):
            continue
        try:
            for name in os.listdir(root):
                if "autocad" in name.lower():
                    cand = os.path.join(root, name, "accoreconsole.exe")
                    if os.path.exists(cand):
                        seen.append(cand)
        except OSError:
            continue
    return sorted(seen)[-1] if seen else ""

def find_converter():
    """(движок, путь): ('acad', ...) | ('oda', ...) | ('', '')."""
    p = find_accore()
    if p:
        return "acad", p
    p = find_oda()
    if p:
        return "oda", p
    return "", ""

def _convert_accore(dwg_paths, exe, out_dir, log):
    """DWG -> DXF через AutoCAD Core Console (по файлу за раз)."""
    result = {}
    tmp = tempfile.mkdtemp(prefix="perechni_acc_")
    for i, p in enumerate(dwg_paths):
        src = os.path.join(tmp, f"in{i}.dwg")     # ASCII-пути: accoreconsole
        dst = os.path.join(tmp, f"out{i}.dxf")    # капризен к кириллице в .scr
        shutil.copy2(p, src)
        scr = os.path.join(tmp, f"c{i}.scr")
        with open(scr, "w", encoding="ascii") as f:
            f.write("_.DXFOUT\n%s\n16\n\n" % dst.replace("\\", "/"))
        log(f"  AutoCAD: {os.path.basename(p)} ...")
        # без pipe (иначе accoreconsole зависает); сам он тоже не всегда
        # завершается — ждём появления DXF и стабилизации его размера
        proc = subprocess.Popen([exe, "/i", src, "/s", scr, "/l", "en-US"],
                                creationflags=getattr(subprocess, "CREATE_NO_WINDOW", 0))
        import time
        stable, last = 0, -1
        for _ in range(300):
            if proc.poll() is not None:
                break
            sz = os.path.getsize(dst) if os.path.exists(dst) else -1
            if sz > 0 and sz == last:
                stable += 1
                if stable >= 3:      # размер не меняется 3 сек — готово
                    break
            else:
                stable = 0
            last = sz
            time.sleep(1)
        if proc.poll() is None:
            proc.kill()
        if not os.path.exists(dst):
            log(f"    ОШИБКА: {os.path.basename(p)} не сконвертировался")
            continue
        if os.path.exists(dst):
            final = os.path.join(out_dir, os.path.splitext(os.path.basename(p))[0] + ".dxf")
            shutil.move(dst, final)
            result[p] = final
            log(f"    готово: {os.path.basename(final)}")
        else:
            log(f"    ОШИБКА: {os.path.basename(p)} не сконвертировался")
    return result

def _convert_oda(dwg_paths, exe, out_dir, log):
    """DWG -> DXF через ODA File Converter (папка за раз)."""
    import time
    tmp_in = tempfile.mkdtemp(prefix="perechni_in_")
    tmp_out = tempfile.mkdtemp(prefix="perechni_out_")
    for p in dwg_paths:
        shutil.copy2(p, os.path.join(tmp_in, os.path.basename(p)))
    proc = subprocess.Popen([exe, tmp_in, tmp_out, "ACAD2018", "DXF", "0", "0", "*.dwg"])
    for _ in range(300):
        time.sleep(1)
        done = [f for f in os.listdir(tmp_out) if f.lower().endswith(".dxf")]
        if len(done) >= len(dwg_paths):
            break
    time.sleep(2)
    try:
        proc.kill()
    except Exception:
        pass
    result = {}
    for p in dwg_paths:
        dxf = os.path.join(tmp_out, os.path.splitext(os.path.basename(p))[0] + ".dxf")
        if os.path.exists(dxf):
            final = os.path.join(out_dir, os.path.basename(dxf))
            shutil.move(dxf, final)
            result[p] = final
            log(f"  готово: {os.path.basename(final)}")
        else:
            log(f"  ОШИБКА: не сконвертировался {os.path.basename(p)}")
    return result

def dwg_to_dxf(dwg_paths, out_dir=None, engine=None, exe=None, log=print):
    """Конвертирует DWG в DXF. Возвращает {dwg: dxf}.
    engine/exe можно не указывать — найдутся сами (AutoCAD, затем ODA)."""
    if not engine or not exe:
        engine, exe = find_converter()
    if not exe:
        raise RuntimeError(
            "Не найден конвертер DWG.\nНужен AutoCAD (accoreconsole.exe) или бесплатный "
            "ODA File Converter (opendesign.com).\nЛибо сохраните чертёж из AutoCAD как DXF.")
    out_dir = out_dir or tempfile.mkdtemp(prefix="perechni_dxf_")
    os.makedirs(out_dir, exist_ok=True)
    log(f"Конвертация {len(dwg_paths)} DWG -> DXF ({'AutoCAD' if engine=='acad' else 'ODA'})...")
    if engine == "acad":
        return _convert_accore(dwg_paths, exe, out_dir, log)
    return _convert_oda(dwg_paths, exe, out_dir, log)

# --------------------------------------------------- показать канал в AutoCAD

def find_acad():
    """Ищет acad.exe (рядом с accoreconsole и в стандартных местах)."""
    import glob as _g
    pats = [r"E:\\AutoCAD*\\acad.exe", r"E:\\autocad*\\acad.exe",
            r"C:\\Program Files\\Autodesk\\AutoCAD*\\acad.exe"]
    for pat in pats:
        hits = sorted(_g.glob(pat), reverse=True)
        if hits:
            return hits[0]
    try:
        eng, exe = find_converter()
        if exe and "accoreconsole" in exe.lower():
            cand = os.path.join(os.path.dirname(exe), "acad.exe")
            if os.path.exists(cand):
                return cand
    except Exception:
        pass
    return ""

def show_in_acad(drawing, x, y, log=print, half_w=130, half_h=90):
    """Открывает чертёж в AutoCAD и зумит на окно вокруг (x, y)."""
    acad = find_acad()
    if not acad:
        raise RuntimeError("Не найден acad.exe — AutoCAD не установлен?")
    if x is None or y is None:
        raise RuntimeError("Для этой строки нет координат (источник — Excel?)")
    scr_dir = os.path.join(tempfile.gettempdir(), "perechni_scr")
    os.makedirs(scr_dir, exist_ok=True)
    scr = os.path.join(scr_dir, "zoom.scr")
    with open(scr, "w", encoding="ascii") as f:
        f.write(f"_.ZOOM\n_W\n{x - half_w:.1f},{y - half_h:.1f}\n"
                f"{x + half_w:.1f},{y + half_h:.1f}\n")
    subprocess.Popen([acad, drawing, "/b", scr])
    log(f"AutoCAD: {os.path.basename(drawing)} -> зум на {x:.0f},{y:.0f}")

# ------------------------------------------------- обновление полей через Word

def update_fields_word(paths, log=print, pdf=False):
    """Открывает Word в фоне (COM), обновляет поля; опционально сохраняет PDF рядом."""
    ps_tpl = r'''
$ErrorActionPreference = "Stop"
$w = New-Object -ComObject Word.Application
$w.Visible = $false
$w.DisplayAlerts = 0
try {{
  $d = $w.Documents.Open("{path}", $false, $false)
  $d.Repaginate()
  $null = $d.Fields.Update()
  foreach ($rng in $d.StoryRanges) {{ $null = $rng.Fields.Update() }}
  $d.Save()
  if ("{pdf}" -ne "") {{ $d.ExportAsFixedFormat("{pdf}", 17) }}
  $d.Close()
}} finally {{ $w.Quit() }}
'''
    ok = True
    for p in paths:
        try:
            pdf_path = (os.path.splitext(p)[0] + ".pdf") if pdf else ""
            r = subprocess.run(
                ["powershell", "-NoProfile", "-Command",
                 ps_tpl.format(path=p.replace('"', ''), pdf=pdf_path.replace('"', ''))],
                capture_output=True, timeout=300)
            if r.returncode == 0:
                log(f"  поля обновлены: {os.path.basename(p)}"
                    + (f" (+PDF)" if pdf else ""))
            else:
                ok = False
                log(f"  не удалось обновить поля: {os.path.basename(p)} "
                    f"(откройте в Word и нажмите «Да»)")
        except Exception:
            ok = False
            log(f"  Word не найден — откройте {os.path.basename(p)} и ответьте «Да» на обновление полей")
    return ok

# ---------------------------------------------------------------- извлечение

POS = re.compile(r"поз\.\s*([^.]+?)\.")
NCH = {"AI": 4, "AO": 4, "WI": 1}

TAG_FALLBACK = re.compile(r"\b[A-Z]{2,4}[0-9]*(?:[-/][A-Z0-9]{1,6})*(?:\(\w+\))?\b")

def _tag(s):
    m = POS.search(s)
    if m:
        return m.group(1).strip()
    # описания без «поз.»: берём первый латинский тег с цифрой (YP9, TIT1, PIT-1(23))
    for m in TAG_FALLBACK.finditer(s):
        t = m.group(0)
        if any(c.isdigit() for c in t):
            return t
    return ""

def _clean(s):
    s = re.sub(r"\s+", " ", s.replace("\n", " ")).strip()
    s = s.replace("авайриного", "аварийного").replace("авариного", "аварийного")
    s = re.sub(r"\s*поз\.\s*[^.]*?\.\s*", " ", s, count=1)
    return re.sub(r"\s{2,}", " ", s).strip()

def extract(dxf_path, log=print):
    """Возвращает список каналов: dict(mod,type,io,ch,tag,desc,level,ctrl,ex,kc)."""
    doc = ezdxf.readfile(dxf_path)
    msp = doc.modelspace()
    T = []
    for e in msp:
        t = e.dxftype()
        if t in ("TEXT", "MTEXT"):
            try:
                s = e.plain_text() if t == "MTEXT" else e.dxf.text
            except Exception:
                s = getattr(e.dxf, "text", "")
            i = e.dxf.insert
            T.append((float(i[0]), float(i[1]), e.dxf.layer, s.strip()))
    hdr = re.compile(r"^Модуль\s+(\d+)\.(\d+)\s+([A-ZА-Я]+)")
    heads = []
    for x, y, l, s in T:
        m = hdr.match(s)
        if m:
            heads.append((x, int(m.group(1)), int(m.group(2)), m.group(3)))
    heads.sort()
    def _is_desc(txt):
        if "поз." in txt:
            return True
        t = txt.strip()
        if len(t) < 20:
            return False
        if t.startswith(("Модуль", "Примечание", "Шкаф", "Питание", "Таблица",
                         "Вид ", "Обозначение")):
            return False
        # два и более кириллических слова = похоже на описание сигнала
        return bool(re.search(r"[А-Яа-яЁё]{4,}[^А-Яа-яЁё]+[А-Яа-яЁё]{3,}", t))
    descs = [(x, y, s) for x, y, l, s in T if l == "Текст" and _is_desc(s)]
    biz = [(x, y) for x, y, l, s in T if l == "mark2" and "БИЗ" in s]
    rows = []
    # аналоговые модули (по заголовкам)
    for i, (x, g, num, typ) in enumerate(heads):
        x1 = heads[i + 1][0] if i + 1 < len(heads) else x + 230
        band = sorted([d for d in descs if x - 30 <= d[0] < x1 - 30], key=lambda r: -r[1])
        has_biz = any(x - 30 <= bx < x1 - 30 for bx, by in biz)
        io_ = "in" if typ in ("AI", "WI") else "out"
        n = max(NCH.get(typ, 4), len(band))
        for ch in range(1, n + 1):
            if ch <= len(band):
                s = band[ch - 1][2]
                tag, desc = _tag(s), _clean(s)
                cx, cy = band[ch - 1][0], band[ch - 1][1]
            else:
                tag, desc = "", "Резерв"
                cx = x
                cy = band[0][1] if band else 3150
            rows.append(dict(sortx=x, x=cx, y=cy, mod=f"{g}.{num}", type=typ, io=io_, ch=ch,
                             tag=tag, desc=desc, level="4-20 мА", ctrl="",
                             ex="Exia" if has_biz else "Exd", kc=f"{g}.{num}.{ch}"))
    # дискретные модули (по якорям KLDI/KLDO)
    for lay, kind, io_ in (("Реле DI (1-KL)", "DI", "in"), ("Реле DO (3-KL)", "DO", "out")):
        pat = re.compile(r"(\d+)\.(\d+)-KLD[IO](\d+)$")
        an = [(x, y, s, pat.match(s)) for x, y, l, s in T if l == lay and pat.match(s)]
        bymod = collections.defaultdict(list)
        for a in an:
            bymod[(int(a[3].group(1)), int(a[3].group(2)))].append(a)
        for (g, num), anch in bymod.items():
            ax = statistics.median(a[0] for a in anch)
            col = [d for d in descs if ax < d[0] < ax + 460]
            colx = min((d[0] for d in col), default=None)
            col = [d for d in col if colx is not None and abs(d[0] - colx) < 10]
            used = set()
            for x, y, s, m in sorted(anch, key=lambda a: -a[1]):
                ch = int(m.group(3))
                best, bd = None, 1e9
                for j, (dx, dy, ds) in enumerate(col):
                    if j in used:
                        continue
                    d = abs(dy - y)
                    if d < bd:
                        bd, best = d, j
                if best is not None and bd < 9:
                    used.add(best)
                    ds = col[best][2]
                    tag, desc = _tag(ds), _clean(ds)
                else:
                    tag, desc = "", "Резерв"
                rows.append(dict(sortx=x, x=x, y=y, mod=f"{g}.{num}", type=kind, io=io_, ch=ch,
                                 tag=tag, desc=desc, level="=24В", ctrl="СК НО",
                                 ex="Exd", kc=f"{g}.{num}.{ch}"))
    rows.sort(key=lambda r: (r["sortx"], r["ch"]))
    n_res = sum(1 for r in rows if r["desc"] == "Резерв")
    if not rows:
        log("  ⚠ КАНАЛЫ НЕ НАЙДЕНЫ — файл не похож на схему подключения.")
        log("    Если это «Чертёж общего вида» (габариты шкафа) — возьмите файл")
        log("    «Схема принципиальная электрическая питания и управления» этого шкафа.")
    else:
        log(f"  извлечено каналов: {len(rows)} (из них резерв: {n_res})")
        if rows and n_res > len(rows) * 0.7:
            log(f"  ⚠ резерва подозрительно много ({n_res}/{len(rows)}) — проверьте раздел на «Сигналах»")
    return rows

# ---------------------------------------------------------------- сборка docx

def _ptext(p):
    return "".join(t.text or "" for t in p.iter(qn("w:t")))

def _para_set_text(p, text):
    rs = p.findall(qn("w:r"))
    if rs:
        first = rs[0]
        for t in first.findall(qn("w:t")):
            first.remove(t)
        for br in first.findall(qn("w:br")):
            first.remove(br)
        for extra in rs[1:]:
            p.remove(extra)
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        first.append(t)
    else:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
        r.append(t); p.append(r)
    for hl in p.findall(qn("w:hyperlink")):
        p.remove(hl)

def _tc_set_text(tc, text):
    for extra in tc.findall(qn("w:p"))[1:]:
        tc.remove(extra)
    _para_set_text(tc.find(qn("w:p")), text)

def _cells(tr):
    return tr.findall(qn("w:tc"))

def _set_title_row(tbl, title):
    tr = tbl.findall(qn("w:tr"))[0]
    seen = set()
    for tc in _cells(tr):
        if id(tc) in seen:
            continue
        seen.add(id(tc)); _tc_set_text(tc, title)

def _build_table(proto, title, data):
    t = copy.deepcopy(proto)
    trs = t.findall(qn("w:tr"))
    proto_row = copy.deepcopy(trs[2])
    for tr in trs[2:]:
        t.remove(tr)
    _set_title_row(t, title)
    for d in data:
        tr = copy.deepcopy(proto_row)
        vals = ["", d["tag"], d["desc"], d["ctrl"], d["level"], d["ex"], d["kc"]]
        for tc, v in zip(_cells(tr), vals):
            _tc_set_text(tc, v)
        t.append(tr)
    return t

_bm = [100]

def _add_bookmark(p, name):
    _bm[0] += 1
    i = str(_bm[0])
    bs = OxmlElement("w:bookmarkStart"); bs.set(qn("w:id"), i); bs.set(qn("w:name"), name)
    be = OxmlElement("w:bookmarkEnd"); be.set(qn("w:id"), i)
    pPr = p.find(qn("w:pPr"))
    if pPr is not None:
        pPr.addnext(bs)
    else:
        p.insert(0, bs)
    p.append(be)

def _cell_pageref(tc, name):
    for extra in tc.findall(qn("w:p"))[1:]:
        tc.remove(extra)
    p = tc.find(qn("w:p"))
    for r in p.findall(qn("w:r")):
        p.remove(r)
    for fs in p.findall(qn("w:fldSimple")):
        p.remove(fs)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), f" PAGEREF {name} \\h ")
    fld.set(qn("w:dirty"), "true")
    r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "0"
    r.append(t); fld.append(r)
    p.append(fld)

def _enable_update_fields(doc):
    s = doc.settings.element
    if s.find(qn("w:updateFields")) is None:
        uf = OxmlElement("w:updateFields"); uf.set(qn("w:val"), "true")
        s.insert(0, uf)

def _fix_date_font(path):
    """Даты в штампе набраны без явного шрифта -> задать ISOCPEUR."""
    import zipfile
    tmp = path + ".tmp"
    OLD = '<w:rFonts w:cs="Arial"/>'
    NEW = '<w:rFonts w:ascii="ISOCPEUR" w:hAnsi="ISOCPEUR" w:cs="Arial"/>'
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for it in zin.infolist():
            data = zin.read(it.filename)
            if it.filename.startswith("word/") and ("header" in it.filename or "footer" in it.filename) \
               and it.filename.endswith(".xml"):
                data = data.decode("utf-8").replace(OLD, NEW).encode("utf-8")
            zout.writestr(it, data)
    os.replace(tmp, path)

def build_docx(template, outfile, io_, word, sections, log=print):
    """
    template  - путь к шаблону-перечню (фирменный вид);
    io_       - 'in' | 'out';
    word      - 'ВХОДНЫЕ' | 'ВЫХОДНЫЕ';
    sections  - список (имя_раздела, rows) в нужном порядке, напр. ("ШКАФ №1", [...]).
    """
    analog_types = ("AI", "WI") if io_ == "in" else ("AO",)
    doc = Document(template)
    body = doc.element.body
    tables = doc.tables
    if len(tables) < 4:
        raise RuntimeError("Шаблон не похож на перечень: мало таблиц")
    toc_tbl = tables[0]._tbl
    analog_proto = copy.deepcopy(tables[1]._tbl)
    discrete_proto = copy.deepcopy(tables[2]._tbl)
    reg_tbl = tables[-1]._tbl
    conts = sectbreak = reg_head = pb_proto = sec_proto = sub_proto = None
    for ch in list(body):
        if ch.tag != qn("w:p"):
            continue
        txt = _ptext(ch).strip()
        pPr = ch.find(qn("w:pPr"))
        if conts is None and "СОДЕРЖАНИЕ" in txt.upper():
            conts = ch
        if sectbreak is None and pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            sectbreak = ch
        if reg_head is None and "ЛИСТ РЕГИСТРАЦИИ" in txt.upper():
            reg_head = ch
        if pb_proto is None and any(b.get(qn("w:type")) == "page" for b in ch.iter(qn("w:br"))):
            pb_proto = ch
        if sub_proto is None and re.match(r"^\d+\.\d+[\s\xa0]+\S", txt):
            sub_proto = ch
        if sec_proto is None and re.match(r"^\d+[\s\xa0]+\S", txt) and not re.match(r"^\d+\.\d", txt):
            sec_proto = ch
    if sec_proto is None:
        sec_proto = sub_proto
    for name, el in (("СОДЕРЖАНИЕ", conts), ("разрыв секции", sectbreak),
                     ("ЛИСТ РЕГИСТРАЦИИ", reg_head), ("разрыв страницы", pb_proto),
                     ("заголовок раздела", sec_proto), ("подзаголовок", sub_proto)):
        if el is None:
            raise RuntimeError(f"В шаблоне не найден элемент: {name}")
    body_sectPr = body.find(qn("w:sectPr"))

    def head(proto, text, bm):
        p = copy.deepcopy(proto)
        _para_set_text(p, text)
        _add_bookmark(p, bm)
        return p

    def pb():
        return copy.deepcopy(pb_proto)

    # содержание
    toc = []
    for i, (cab, _r) in enumerate(sections, 1):
        toc.append((f"{i} {cab.upper()}", f"bmsec{i}"))
        toc.append((f"{i}.1 Аналоговые {word.lower()} сигналы. {cab.capitalize()}", f"bm{i}_1"))
        toc.append((f"{i}.2 Дискретные {word.lower()} сигналы. {cab.capitalize()}", f"bm{i}_2"))
    toc.append(("ЛИСТ РЕГИСТРАЦИИ ИЗМЕНЕНИЙ", "bmreg"))
    trs = toc_tbl.findall(qn("w:tr"))
    proto = copy.deepcopy(trs[0])
    for tr in trs:
        toc_tbl.remove(tr)
    for text, bm in toc:
        tr = copy.deepcopy(proto)
        cs = _cells(tr)
        _tc_set_text(cs[0], text)
        if len(cs) > 1:
            _cell_pageref(cs[1], bm)
        toc_tbl.append(tr)

    _add_bookmark(reg_head, "bmreg")
    new = [conts, toc_tbl, sectbreak]
    for i, (cab, rows) in enumerate(sections, 1):
        an = [r for r in rows if r["io"] == io_ and r["type"] in analog_types]
        di = [r for r in rows if r["io"] == io_ and r["type"] in ("DI", "DO")]
        if i > 1:
            new.append(pb())
        new.append(head(sec_proto, f"{i} {cab.upper()}", f"bmsec{i}"))
        new.append(head(sub_proto, f"{i}.1 АНАЛОГОВЫЕ {word} СИГНАЛЫ. {cab.upper()}", f"bm{i}_1"))
        new.append(_build_table(analog_proto,
                   f"Таблица {i}.1 Аналоговые {word.lower()} сигналы. {cab.capitalize()}", an))
        new.append(pb())
        new.append(head(sub_proto, f"{i}.2 ДИСКРЕТНЫЕ {word} СИГНАЛЫ. {cab.upper()}", f"bm{i}_2"))
        new.append(_build_table(discrete_proto,
                   f"Таблица {i}.2 Дискретные {word.lower()} сигналы. {cab.capitalize()}", di))
        log(f"  {cab}: аналоговых {len(an)}, дискретных {len(di)}")
    new.append(pb())
    new.append(reg_head)
    new.append(reg_tbl)
    new.append(body_sectPr)
    for ch in list(body):
        body.remove(ch)
    for e in new:
        body.append(e)
    _enable_update_fields(doc)
    doc.save(outfile)
    _fix_date_font(outfile)
    log(f"  сохранено: {outfile}")

# ---------------------------------------------------------------- Excel-сводка

def export_xlsx(sections, path):
    """Сводка сигналов в .xlsx: лист «Сигналы» + лист «Сводка».
    sections: [(имя_раздела, rows)]"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    thin = Border(*[Side(style="thin", color="D0D7E2")] * 4)
    head_fill = PatternFill("solid", fgColor="2563EB")
    head_font = Font(color="FFFFFF", bold=True, size=10)
    alt_fill = PatternFill("solid", fgColor="F3F6FB")
    res_font = Font(color="9CA3AF", italic=True)

    ws = wb.active
    ws.title = "Сигналы"
    cols = ["Шкаф", "Позиция в контроллере", "Тип", "Позиция по проекту",
            "Описание сигнала", "Уровень управления", "Уровень сигнала", "Взрывозащита"]
    ws.append(cols)
    for c in ws[1]:
        c.fill, c.font, c.border = head_fill, head_font, thin
        c.alignment = Alignment(horizontal="center", vertical="center")
    r = 2
    for cab, rows in sections:
        for d in rows:
            ws.append([cab, d["kc"], d["type"], d["tag"], d["desc"],
                       d["ctrl"], d["level"], d["ex"]])
            for c in ws[r]:
                c.border = thin
                if r % 2 == 0:
                    c.fill = alt_fill
                if d["desc"] == "Резерв":
                    c.font = res_font
            r += 1
    widths = [12, 20, 7, 18, 70, 12, 12, 12]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"
    ws.auto_filter.ref = f"A1:H{r - 1}"

    sm = wb.create_sheet("Сводка")
    sm.append(["Шкаф", "Тип", "Всего каналов", "Занято", "Резерв"])
    for c in sm[1]:
        c.fill, c.font, c.border = head_fill, head_font, thin
    r = 2
    for cab, rows in sections:
        by = {}
        for d in rows:
            t = by.setdefault(d["type"], [0, 0])
            t[0] += 1
            if d["desc"] != "Резерв":
                t[1] += 1
        for typ in sorted(by):
            tot, used = by[typ]
            sm.append([cab, typ, tot, used, tot - used])
            for c in sm[r]:
                c.border = thin
            r += 1
    for i, w in enumerate([14, 8, 14, 10, 10], 1):
        sm.column_dimensions[get_column_letter(i)].width = w
    wb.save(path)
    return path

# ------------------------------------------------- разделы из Excel-сводки

def sections_from_xlsx(path, log=print):
    """Читает Excel-сводку (лист «Сигналы»: Шкаф|Тип|Позиция по проекту|Описание|
    Уровень управления|Уровень сигнала|Взрывозащита|Позиция в контроллере).
    Возвращает [(имя_раздела, rows)] — по одному разделу на каждый «Шкаф»."""
    from openpyxl import load_workbook
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb["Сигналы"] if "Сигналы" in wb.sheetnames else wb.active
    rows_iter = ws.iter_rows(values_only=True)
    header = None
    data = []
    for row in rows_iter:
        vals = ["" if v is None else str(v).strip() for v in row]
        if header is None:
            if any("Позиция" in v for v in vals) and any("Описание" in v for v in vals):
                header = [v.lower() for v in vals]
            continue
        if not any(vals):
            continue
        data.append(vals)
    if header is None:
        raise RuntimeError(f"В {os.path.basename(path)} не найден заголовок таблицы сигналов")

    def col(*keys):
        for i, h in enumerate(header):
            if any(k in h for k in keys):
                return i
        return None
    c_cab, c_typ = col("шкаф", "раздел"), col("тип")
    c_tag, c_desc = col("позиция по проекту", "тег"), col("описание")
    c_ctrl, c_lvl = col("уровень управления"), col("уровень сигнала")
    c_ex, c_kc = col("взрыв"), col("позиция в конт")
    if c_desc is None or c_kc is None:
        raise RuntimeError("Не хватает колонок «Описание» / «Позиция в контроллере»")
    secs, order = {}, []
    for v in data:
        def g(ci, default=""):
            return v[ci] if ci is not None and ci < len(v) else default
        cab = g(c_cab) or cab_name_from_file(path)
        typ = g(c_typ).upper() or ("DI" if "24" in g(c_lvl) else "AI")
        io_ = "in" if typ in ("AI", "WI", "DI") else "out"
        kc = g(c_kc)
        m = re.match(r"(\d+)\.(\d+)\.(\d+)", kc)
        rows = secs.setdefault(cab, [])
        if cab not in order:
            order.append(cab)
        rows.append(dict(sortx=len(rows), mod=(f"{m.group(1)}.{m.group(2)}" if m else ""),
                         type=typ, io=io_, ch=(int(m.group(3)) if m else len(rows) + 1),
                         tag=g(c_tag), desc=g(c_desc) or "Резерв", level=g(c_lvl),
                         ctrl=g(c_ctrl), ex=g(c_ex), kc=kc))
    log(f"  из Excel: {sum(len(r) for r in secs.values())} каналов, разделов: {len(order)}")
    return [(cab, secs[cab]) for cab in order]

# --------------------------------------------------------- свободные каналы

def free_channels(sections):
    """[(шкаф, модуль, тип, свободно, 'позиции')] по строкам «Резерв»."""
    out = []
    for cab, rows in sections:
        bymod = {}
        for d in rows:
            if d.get("desc") == "Резерв" and d.get("mod"):
                bymod.setdefault((d["mod"], d["type"]), []).append(d["kc"])
        for (mod, typ), kcs in sorted(bymod.items(),
                                      key=lambda k: [int(v) for v in k[0][0].split(".")]):
            out.append((cab, mod, typ, len(kcs), ", ".join(kcs)))
    return out

# ------------------------------------------------------ проверки нормоконтроля

def checks_report(sections):
    """Проверки данных перечней. Возвращает список предупреждений (строк)."""
    warns = []
    for cab, rows in sections:
        seen_kc = {}
        tag_desc = {}
        for d in rows:
            kc = d.get("kc", "")
            if kc:
                if kc in seen_kc:
                    warns.append(f"{cab}: дубль позиции {kc} (каналы задвоены)")
                seen_kc[kc] = True
            tag, desc = d.get("tag", ""), d.get("desc", "")
            if desc == "Резерв":
                if tag:
                    warns.append(f"{cab}: {kc} — «Резерв» с непустым тегом «{tag}»")
                continue
            if tag and not desc:
                warns.append(f"{cab}: {kc} — тег «{tag}» без описания")
            if desc and not tag:
                warns.append(f"{cab}: {kc} — описание без тега (позиции по проекту)")
            if tag:
                tag_desc.setdefault(tag, desc)
        # непрерывность каналов внутри модулей
        bymod = {}
        for d in rows:
            if d.get("mod"):
                bymod.setdefault((d["mod"], d["type"]), []).append(d["ch"])
        for (mod, typ), chs in bymod.items():
            chs = sorted(chs)
            missing = [c for c in range(1, max(chs) + 1) if c not in chs]
            if missing:
                warns.append(f"{cab}: модуль {mod} {typ} — пропущены каналы {missing}")
    # один тег — разные описания (между шкафами тоже)
    tagmap = {}
    for cab, rows in sections:
        for d in rows:
            if d.get("tag") and d.get("desc") != "Резерв":
                tagmap.setdefault((d["tag"], d.get("type")), set()).add(d["desc"])
    for (tag, typ), descs in tagmap.items():
        if len(descs) > 2:
            warns.append(f"Тег {tag} ({typ}) имеет {len(descs)} разных описаний — проверьте")
    return warns

# --------------------------------------------------------- спецификация и КИП

_SKIP_BLOCKS = ("штамп", "рамка", "sw_", "формат", "solid", "wipeout", "подпись")

def extract_equipment(dxf_path):
    """Состав шкафа из чертежа: модули (по заголовкам), блоки/клеммы (по вставкам).
    Возвращает список (элемент, артикул, кол-во, примечание)."""
    import collections
    doc = ezdxf.readfile(dxf_path)
    msp = doc.modelspace()
    texts, inserts = [], collections.Counter()
    for e in msp:
        t = e.dxftype()
        if t in ("TEXT", "MTEXT"):
            try:
                s = e.plain_text() if t == "MTEXT" else e.dxf.text
            except Exception:
                s = getattr(e.dxf, "text", "")
            texts.append((e.dxf.layer, s.strip()))
        elif t == "INSERT":
            nm = e.dxf.name
            if nm and not any(k in nm.lower() for k in _SKIP_BLOCKS):
                inserts[nm] += 1
    out = []
    # модули с артикулами из полных заголовков
    hdr = re.compile(r"^Модуль\s+(ввода|вывода)\s+([а-яё\s]+?)\s+сигналов\s+[\d.]+\s+"
                     r"([A-Z]+),\s*([0-9A-Z\-]+)", re.I)
    mods = {}
    for lay, s in texts:
        m = hdr.match(s)
        if m:
            key = (f"Модуль {m.group(1).lower()} {m.group(2).lower()} сигналов {m.group(3)}",
                   m.group(4))
            mods[key] = mods.get(key, 0) + 1
    kinds_in_headers = set()
    for (name, art), n in sorted(mods.items()):
        out.append((name, art, n, "по заголовкам листов"))
        m = re.search(r"\b(AI|AO|DI|DO|WI)\b", name)
        if m:
            kinds_in_headers.add(m.group(1))
    # дискретные модули по якорям каналов (если не посчитаны по заголовкам)
    for lay, kind in (("Реле DI (1-KL)", "DI"), ("Реле DO (3-KL)", "DO")):
        if kind in kinds_in_headers:
            continue
        nums = set()
        pat = re.compile(r"(\d+\.\d+)-KLD[IO]\d+$")
        for l, s in texts:
            if l == lay:
                m = pat.match(s)
                if m:
                    nums.add(m.group(1))
        if nums:
            out.append((f"Модуль дискретных сигналов {kind} 16 каналов", "", len(nums),
                        "по каналам " + ", ".join(sorted(nums))))
    # интерфейсный модуль и прочее из зоны состава
    for l, s in texts:
        m = re.search(r"интерфейсный модуль\s+(\S+.*?)\s*\(([\w\-]+)\)", s, re.I)
        if m:
            out.append(("Интерфейсный модуль " + m.group(1), m.group(2), 1, "из состава на чертеже"))
            break
    # блоки/клеммы по вставкам
    for nm, n in inserts.most_common():
        note = ""
        if "биз" in nm.lower() or "rpssi" in nm.lower():
            note = "барьер искрозащиты"
        elif "degson" in nm.lower() or "клемма" in nm.lower():
            note = "клемма"
        out.append((nm, "", n, note))
    return out

def read_equipment(drawings, oda_exe=None, log=print):
    """Состав по списку чертежей (DWG конвертируются) -> [(имя_раздела, items)]."""
    items = [(d, None) if isinstance(d, str) else (d[0], d[1]) for d in drawings]
    dwgs = [p for p, _n in items if p.lower().endswith(".dwg")]
    conv = {}
    if dwgs:
        engine, exe = ("oda", oda_exe) if oda_exe else find_converter()
        conv = dwg_to_dxf(dwgs, engine=engine, exe=exe, log=log)
    out = []
    for p, name in items:
        dxf = conv.get(p, p)
        log(f"Состав: {os.path.basename(p)}")
        out.append((name or cab_name_from_file(p), extract_equipment(dxf)))
    return out

def export_spec_xlsx(equip_sections, sections, path):
    """equip_sections: [(шкаф, [(элемент, артикул, кол-во, прим)])];
    sections: [(шкаф, rows)] для ведомости КИП."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    wb = Workbook()
    thin = Border(*[Side(style="thin", color="D0D7E2")] * 4)
    hf = PatternFill("solid", fgColor="2563EB")
    hfont = Font(color="FFFFFF", bold=True, size=10)

    ws = wb.active
    ws.title = "Состав шкафов"
    ws.append(["Шкаф", "Элемент", "Артикул", "Кол-во", "Примечание"])
    for c in ws[1]:
        c.fill, c.font, c.border = hf, hfont, thin
    r = 2
    for cab, items in equip_sections:
        for name, art, n, note in items:
            ws.append([cab, name, art, n, note])
            for c in ws[r]:
                c.border = thin
            r += 1
    for i, w in enumerate([14, 58, 22, 8, 40], 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"

    kd = wb.create_sheet("Ведомость КИП")
    kd.append(["Позиция", "Описание (по первому каналу)", "Сигналы", "Каналов", "Шкаф"])
    for c in kd[1]:
        c.fill, c.font, c.border = hf, hfont, thin
    tags = {}
    for cab, rows in sections:
        for d in rows:
            if not d["tag"] or d["desc"] == "Резерв":
                continue
            t = tags.setdefault(d["tag"], dict(desc=d["desc"], types=set(), n=0, cabs=set()))
            t["types"].add(d["type"])
            t["n"] += 1
            t["cabs"].add(cab)
    r = 2
    for tag in sorted(tags):
        t = tags[tag]
        kd.append([tag, t["desc"], "/".join(sorted(t["types"])), t["n"],
                   ", ".join(sorted(t["cabs"]))])
        for c in kd[r]:
            c.border = thin
        r += 2 - 1
    for i, w in enumerate([16, 75, 12, 9, 18], 1):
        kd.column_dimensions[get_column_letter(i)].width = w
    kd.freeze_panes = "A2"
    kd.auto_filter.ref = f"A1:E{r - 1}"
    wb.save(path)
    return path

# ---------------------------------------------------------- штампы (даты и пр.)

_P_RX = re.compile(r"<w:p(?=[ >]).*?</w:p>", re.S)
_T_RX = re.compile(r"(<w:t[^>]*>)(.*?)(</w:t>)", re.S)

def _hf_parts(zf):
    return [n for n in zf.namelist()
            if n.startswith("word/") and n.endswith(".xml")
            and ("header" in n or "footer" in n)]

def detect_stamp_value(path, pattern=r"^\d{2}[.,]\d{2}$"):
    """Самое частое значение-дата в штампах (колонтитулы, включая надписи)."""
    import zipfile, collections
    rx = re.compile(pattern)
    cnt = collections.Counter()
    with zipfile.ZipFile(path) as zf:
        for name in _hf_parts(zf):
            xml = zf.read(name).decode("utf-8")
            for m in _P_RX.finditer(xml):
                txt = "".join(t[1] for t in _T_RX.findall(m.group(0))).strip()
                if rx.match(txt):
                    cnt[txt] += 1
    return cnt.most_common(1)[0][0] if cnt else ""

def replace_in_stamps(paths, old, new, log=print, backup=True):
    """Заменяет значения (напр. дату «05.26») в штампах-колонтитулах docx."""
    import zipfile
    if not old or not str(old).strip():
        raise RuntimeError("Пустое старое значение — нечего заменять")
    total = 0
    for path in paths:
        n = 0
        with zipfile.ZipFile(path) as zf:
            data = {name: zf.read(name) for name in zf.namelist()}
        for name in list(data):
            if not (name.startswith("word/") and name.endswith(".xml")
                    and ("header" in name or "footer" in name)):
                continue
            xml = data[name].decode("utf-8")

            def fix_para(m):
                nonlocal n
                frag = m.group(0)
                txt = "".join(t[1] for t in _T_RX.findall(frag)).strip()
                if txt != old:
                    return frag
                first = [True]

                def sub_t(mt):
                    val = new if first[0] else ""
                    first[0] = False
                    return mt.group(1) + val + mt.group(3)
                n += 1
                return _T_RX.sub(sub_t, frag)
            data[name] = _P_RX.sub(fix_para, xml).encode("utf-8")
        if n:
            if backup:
                shutil.copy2(path, path + ".bak")
            with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
                for name, blob in data.items():
                    zout.writestr(name, blob)
        log(f"  {os.path.basename(path)}: замен {n}")
        total += n
    return total

# ------------------------------------------------------------- заготовка ПЗ

def export_pz_docx(sections, equip_sections, path):
    """Заготовка раздела пояснительной записки: сводка каналов, состав ТС, КИП."""
    doc = Document()
    doc.add_heading("Автоматизация. Сигналы и технические средства (заготовка)", 1)
    doc.add_paragraph(
        "Раздел сформирован автоматически по чертежам схем подключения. "
        "Текст подлежит редактированию и нормоконтролю.")
    doc.add_heading("1 Сводка сигналов", 2)
    total = dict()
    for cab, rows in sections:
        by = {}
        for d in rows:
            t = by.setdefault(d["type"], [0, 0])
            t[0] += 1
            if d["desc"] != "Резерв":
                t[1] += 1
            g = total.setdefault(d["type"], [0, 0])
            g[0] += 1
            if d["desc"] != "Резерв":
                g[1] += 1
        parts = [f"{k}: {v[1]} из {v[0]} (резерв {v[0]-v[1]})" for k, v in sorted(by.items())]
        doc.add_paragraph(f"{cab}: " + "; ".join(parts) + ".")
    parts = [f"{k}: {v[1]} из {v[0]}" for k, v in sorted(total.items())]
    doc.add_paragraph("Итого по системе: " + "; ".join(parts) + ".")

    if equip_sections:
        doc.add_heading("2 Состав технических средств", 2)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        for i, h in enumerate(("Шкаф", "Элемент", "Артикул", "Кол-во")):
            tbl.rows[0].cells[i].text = h
        for cab, items in equip_sections:
            for name, art, n, _note in items:
                r = tbl.add_row().cells
                r[0].text, r[1].text, r[2].text, r[3].text = cab, name, art, str(n)

    doc.add_heading("3 Ведомость приборов КИП", 2)
    tags = {}
    for cab, rows in sections:
        for d in rows:
            if not d["tag"] or d["desc"] == "Резерв":
                continue
            t = tags.setdefault(d["tag"], dict(desc=d["desc"], types=set(), cabs=set()))
            t["types"].add(d["type"])
            t["cabs"].add(cab)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for i, h in enumerate(("Позиция", "Назначение", "Сигналы", "Шкаф")):
        tbl.rows[0].cells[i].text = h
    for tag in sorted(tags):
        t = tags[tag]
        r = tbl.add_row().cells
        r[0].text, r[1].text = tag, t["desc"]
        r[2].text, r[3].text = "/".join(sorted(t["types"])), ", ".join(sorted(t["cabs"]))
    doc.save(path)
    return path

# ------------------------------------------------------------- сравнение версий

def parse_perechen(path):
    """Читает существующий перечень .docx -> [{cab, kc, tag, desc, ctrl, level, ex}]."""
    doc = Document(path)
    out = []
    for t in doc.tables:
        if len(t.columns) != 7 or len(t.rows) < 3:
            continue
        title = t.rows[0].cells[0].text.strip()
        m = re.search(r"(?:Шкаф|Щит)\s+(\S+)\s*$", title, re.I)
        cab = ("ШКАФ " + m.group(1).upper()) if m else ""
        io_ = "in" if "входн" in title.lower() else ("out" if "выходн" in title.lower() else "")
        for row in t.rows[2:]:
            cs = [c.text.strip() for c in row.cells]
            if len(cs) < 7 or not cs[6]:
                continue
            out.append(dict(cab=cab, io=io_, tag=cs[1], desc=re.sub(r"\s+", " ", cs[2]),
                            ctrl=cs[3], level=cs[4], ex=cs[5], kc=cs[6]))
    return out

def compare(sections, old_rows, log=print):
    """Сравнивает извлечённое из чертежей со старым перечнем.
    Возвращает список (статус, шкаф, kc, было, стало)."""
    new = {}
    for cab, rows in sections:
        for d in rows:
            new[(cab.upper(), d["kc"])] = d
    old = {(o["cab"].upper(), o["kc"]): o for o in old_rows}
    report = []
    for key in sorted(set(new) | set(old), key=lambda k: (k[0], k[1])):
        n, o = new.get(key), old.get(key)
        cab, kc = key
        if n and not o:
            if n["desc"] != "Резерв":
                report.append(("добавлено", cab, kc, "", f"{n['tag']} | {n['desc']}"))
        elif o and not n:
            if o["desc"] != "Резерв":
                report.append(("удалено", cab, kc, f"{o['tag']} | {o['desc']}", ""))
        else:
            was = f"{o['tag']} | {o['desc']}"
            now = f"{n['tag']} | {n['desc']}"
            if (o["tag"] != n["tag"] or o["desc"] != n["desc"]) and not (
                    o["desc"] == "Резерв" and n["desc"] == "Резерв"):
                if o["desc"] == "Резерв" and n["desc"] != "Резерв":
                    report.append(("добавлено", cab, kc, "Резерв", now))
                elif n["desc"] == "Резерв" and o["desc"] != "Резерв":
                    report.append(("удалено", cab, kc, was, "Резерв"))
                else:
                    report.append(("изменено", cab, kc, was, now))
    log(f"Сравнение: изменений {len(report)} "
        f"(добавлено {sum(1 for r in report if r[0]=='добавлено')}, "
        f"удалено {sum(1 for r in report if r[0]=='удалено')}, "
        f"изменено {sum(1 for r in report if r[0]=='изменено')})")
    return report

def export_compare_xlsx(report, path):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    wb = Workbook()
    ws = wb.active
    ws.title = "Изменения"
    thin = Border(*[Side(style="thin", color="D0D7E2")] * 4)
    ws.append(["Статус", "Шкаф", "Позиция", "Было", "Стало"])
    fills = {"добавлено": PatternFill("solid", fgColor="DCFCE7"),
             "удалено": PatternFill("solid", fgColor="FEE2E2"),
             "изменено": PatternFill("solid", fgColor="FEF9C3")}
    for c in ws[1]:
        c.font = Font(bold=True)
        c.border = thin
    for i, (st, cab, kc, was, now) in enumerate(report, 2):
        ws.append([st, cab, kc, was, now])
        for c in ws[i]:
            c.border = thin
            c.fill = fills.get(st, PatternFill())
    for i, w in enumerate([12, 14, 12, 60, 60], 1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"
    wb.save(path)
    return path

def sections_from_perechni(paths, log=print):
    """Готовит данные для генератора схем из перечней .docx (аналоговые каналы)."""
    rows = []
    for p in paths:
        if p:
            rows += parse_perechen(p)
    bycab = {}
    for o in rows:
        if "мА" not in o.get("level", ""):
            continue  # дискретные листы — следующий этап
        m = re.match(r"^(\d+\.\d+)\.(\d+)$", o["kc"].strip())
        if not m:
            continue
        typ = "AI" if o.get("io") == "in" else "AO"
        d = dict(mod=m.group(1), ch=int(m.group(2)), type=typ, io=o.get("io", ""),
                 tag=o["tag"], desc=o["desc"], kc=o["kc"], level=o["level"],
                 ctrl=o.get("ctrl", ""), ex=o.get("ex", ""))
        bycab.setdefault(o["cab"] or "ШКАФ", []).append(d)
    out = []
    for cab in bycab:
        rs = sorted(bycab[cab], key=lambda r: ([int(x) for x in r["mod"].split(".")], r["ch"]))
        out.append((cab, rs))
        log(f"  {cab}: аналоговых каналов {len(rs)}")
    return out

# --------------------------------------------------------- генерация схем (бета)

def _sheet_frames(msp):
    out = []
    for e in msp.query("LWPOLYLINE"):
        pts = e.get_points("xy")
        xs = [p[0] for p in pts]; ys = [p[1] for p in pts]
        w, h = max(xs) - min(xs), max(ys) - min(ys)
        if 180 < w < 240 and 260 < h < 320:
            out.append((min(xs), min(ys), max(xs), max(ys)))
    out.sort()
    return out

def _anchor(e):
    t = e.dxftype()
    try:
        if t in ("TEXT", "MTEXT", "INSERT"):
            p = e.dxf.insert; return p[0], p[1]
        if t == "LINE":
            return (e.dxf.start[0] + e.dxf.end[0]) / 2, (e.dxf.start[1] + e.dxf.end[1]) / 2
        if t in ("CIRCLE", "ARC", "ELLIPSE"):
            c = e.dxf.center; return c[0], c[1]
        if t == "LWPOLYLINE":
            pts = e.get_points("xy")
            return (sum(p[0] for p in pts) / len(pts), sum(p[1] for p in pts) / len(pts))
        if t == "SPLINE":
            cp = list(e.control_points)
            return (sum(p[0] for p in cp) / len(cp), sum(p[1] for p in cp) / len(cp))
    except Exception:
        pass
    return None

def generate_scheme(donor_dxf, out_dxf, sections, types=("AI", "AO"), log=print):
    """Генерирует листы схемы подключения в стиле донора (бета: аналоговые модули).
    Донор — фирменный чертёж, из которого берётся по одному заполненному листу
    каждого типа; на каждый модуль из sections печатается свой лист."""
    from ezdxf.math import Matrix44
    doc = ezdxf.readfile(donor_dxf)
    msp = doc.modelspace()
    frames = _sheet_frames(msp)
    if not frames:
        raise RuntimeError("В доноре не найдены рамки листов А4")
    hdr = re.compile(r"^Модуль\s+(\d+\.\d+)\s+([A-Z]+)")
    kcre = re.compile(r"^(\d+\.\d+\.\d+)([+-])$")

    def sheet_of(x, y):
        for f in frames:
            if f[0] <= x <= f[2] and f[1] <= y <= f[3]:
                return f
        return None

    # донорские листы по типам: первый лист типа с максимумом описаний
    donors = {}
    for t in msp.query("TEXT"):
        m = hdr.match(t.dxf.text.strip())
        if not m or m.group(2) not in types:
            continue
        f = sheet_of(t.dxf.insert[0], t.dxf.insert[1])
        if not f:
            continue
        ndesc = sum(1 for e in msp.query("MTEXT")
                    if e.dxf.layer == "Текст" and "поз." in e.text
                    and f[0] <= e.dxf.insert[0] <= f[2] and f[1] <= e.dxf.insert[1] <= f[3])
        cur = donors.get(m.group(2))
        if cur is None or ndesc > cur[2]:
            donors[m.group(2)] = (f, m.group(1), ndesc)
    PAD = 12
    all_x = [a[0] for e in msp if (a := _anchor(e))]
    gen_x = max(all_x) + 500
    sheet_no = [1]
    biz_no = [1]
    new_ids = set()

    def build_donor(typ):
        f, dmod, _n = donors[typ]
        x0, y0, x1, y1 = f
        ents = [e for e in msp if (a := _anchor(e)) and
                x0 - PAD <= a[0] <= x1 + PAD and y0 - 30 <= a[1] <= y1 + PAD]
        # каналы донора: КС-метки с +/-
        kc = {}
        for e in ents:
            if e.dxftype() == "TEXT":
                m = kcre.match(e.dxf.text.strip())
                if m:
                    kc.setdefault(m.group(1), []).append(e.dxf.insert[1])
        chans = sorted(kc, key=lambda k: -max(kc[k]))
        # теги донора: подписи вида ТЕГ+ / ТЕГ- / ТЕГ-U.1 (не КС)
        sufre = re.compile(r"^(.{2,}?)([+-]|-U\.\d+)$")
        tags = {}
        for e in ents:
            if e.dxftype() == "TEXT":
                t = e.dxf.text.strip()
                m = sufre.match(t)
                if m and not kcre.match(t):
                    tags.setdefault(m.group(1).rstrip("-"), []).append(e.dxf.insert[1])
        dtags = sorted(tags, key=lambda k: -max(tags[k]))
        # описания сверху вниз: текст донора -> номер канала
        descs = sorted([e for e in ents if e.dxftype() == "MTEXT"
                        and e.dxf.layer == "Текст" and "поз." in e.text],
                       key=lambda e: -e.dxf.insert[1])
        dmap = {e.text: i for i, e in enumerate(descs)}
        return dict(f=f, mod=dmod, ents=ents, chans=chans, tags=dtags, dmap=dmap)

    dcache = {t: build_donor(t) for t in donors}
    made = 0
    for cab, rows in sections:
        mods = {}
        for r in rows:
            if r["type"] in dcache:
                mods.setdefault((r["type"], r["mod"]), []).append(r)
        for (typ, mod), chans in sorted(
                mods.items(), key=lambda kv: [int(x) for x in kv[0][1].split(".")]):
            d = dcache[typ]
            x0 = d["f"][0]
            dx = gen_x - x0 + 240.0 * made
            repl = {}
            repl[f"Модуль {d['mod']} {typ}"] = f"Модуль {mod} {typ}"
            repl[f"сигналов {d['mod']} {typ}"] = f"сигналов {mod} {typ}"
            for i, dkc in enumerate(d["chans"]):
                nk = chans[i]["kc"] if i < len(chans) else f"{mod}.{i+1}"
                repl[dkc + "+"] = nk + "+"
                repl[dkc + "-"] = nk + "-"
                repl[dkc] = nk
            tag_map = {}
            for i, dt in enumerate(d["tags"]):
                ch = chans[i] if i < len(chans) else None
                tag_map[dt] = (ch["tag"] if ch and ch["desc"] != "Резерв" else "")
            new = []
            for e in d["ents"]:
                c = e.copy()
                c.transform(Matrix44.translate(dx, 0, 0))
                msp.add_entity(c)
                new.append(c)
                new_ids.add(id(c))
            for e in new:
                if e.dxftype() == "TEXT":
                    t = e.dxf.text.strip()
                    for old, nv in repl.items():
                        if t == old:
                            e.dxf.text = nv
                            break
                    else:
                        hit = None
                        for dt in sorted(tag_map, key=len, reverse=True):
                            if t == dt or (t.startswith(dt) and
                                           re.match(r"^([+-]|-U\.\d+)$", t[len(dt):])):
                                hit = dt
                                break
                        if hit is not None:
                            nt = tag_map[hit]
                            e.dxf.text = (nt + t[len(hit):]) if nt else ""
                        elif "БИЗ-" in t and e.dxf.layer == "mark2":
                            e.dxf.text = f"БИЗ-{typ}-{biz_no[0]}"
                            biz_no[0] += 1
                        else:
                            for old, nv in repl.items():
                                if old in t:
                                    e.dxf.text = t.replace(old, nv)
                                    break
                elif e.dxftype() == "MTEXT":
                    t = e.text
                    if e.dxf.layer == "Текст" and t in d["dmap"]:
                        i = d["dmap"][t]
                        ch = chans[i] if i < len(chans) else None
                        if ch is None or ch["desc"] == "Резерв" or not ch["tag"]:
                            e.text = "Резерв"
                        else:
                            e.text = f"{ch['desc']}. Поз. {ch['tag']}"
                    elif "БИЗ-" in t and "См." in t:
                        e.text = re.sub(r"БИЗ-\w+-\d+\s*", f"БИЗ-{typ}-{biz_no[0]} ", t, count=1)
                    elif e.dxf.layer.startswith("Num_str") and t.strip().isdigit():
                        e.text = str(sheet_no[0])
                    else:
                        for old, nv in repl.items():
                            if old in t:
                                e.text = t.replace(old, nv)
                                break
            made += 1
            sheet_no[0] += 1
            log(f"  лист: {cab} Модуль {mod} {typ} ({min(len(chans), len(d['chans']))} кан.)")
    # удалить всё, кроме новых листов
    for e in list(msp):
        if id(e) not in new_ids:
            try:
                msp.delete_entity(e)
            except Exception:
                pass
    doc.saveas(out_dxf)
    log(f"Схема сохранена: {out_dxf} (листов: {made}) — БЕТА: аналоговые модули")
    return out_dxf

# ---------------------------------------------------------------- оркестрация

# С/C и К/K допускаем и кириллицей и латиницей (в именах файлов их путают)
CAB_RE = re.compile(r"(Ш[СC][КK]\s*\d+|Ш[СC]\s*\d+|\d+\s*Щ\b|Щ\s*\d+)", re.I)

def cab_name_from_file(path):
    name = os.path.basename(path)
    m = CAB_RE.search(name)
    if m:
        val = re.sub(r"\s+", "", m.group(1).upper())
        # нормализуем латиницу в кириллицу
        val = val.replace("C", "С").replace("K", "К")
        return ("ЩИТ " if "Щ" in val else "ШКАФ ") + val
    return os.path.splitext(name)[0]

def read_sections(drawings, oda_exe=None, log=print):
    """Читает чертежи (DWG конвертируются) -> [(имя_раздела, rows)].
    drawings: пути или пары (путь, имя_раздела)."""
    items = [(d, None) if isinstance(d, str) else (d[0], d[1]) for d in drawings]
    dwgs = [p for p, _n in items if p.lower().endswith(".dwg")]
    conv = {}
    if dwgs:
        engine, exe = ("oda", oda_exe) if oda_exe else find_converter()
        conv = dwg_to_dxf(dwgs, engine=engine, exe=exe, log=log)
        missing = [p for p in dwgs if p not in conv]
        if missing:
            raise RuntimeError("Не сконвертировались: " + ", ".join(os.path.basename(m) for m in missing))
    sections = []
    for p, name in items:
        log(f"Чтение: {os.path.basename(p)}")
        if p.lower().endswith(".xlsx"):
            for cab, rows in sections_from_xlsx(p, log):
                sections.append((name or cab, rows))
            continue
        dxf = conv.get(p, p)
        rows = extract(dxf, log)
        for r in rows:
            r["src"] = p
        seen = {}
        for r in rows:
            if r["kc"] in seen:
                log(f"  ВНИМАНИЕ: позиция {r['kc']} встречается дважды")
            seen[r["kc"]] = True
        sections.append((name or cab_name_from_file(p), rows))
    return sections

def run(drawings, template_in, template_out, out_dir, oda_exe=None, log=print,
        update_fields=False, sections=None, make_pdf=False):
    """drawings: список путей DWG/DXF или пар (путь, имя_раздела) в порядке разделов.
    sections — готовые данные (например, правленные в предпросмотре); если
    переданы, чертежи не читаются повторно."""
    if sections is None:
        sections = read_sections(drawings, oda_exe=oda_exe, log=log)
    os.makedirs(out_dir, exist_ok=True)
    results = []
    if template_in:
        out = os.path.join(out_dir, os.path.basename(template_in))
        log("Сборка перечня входных сигналов...")
        build_docx(template_in, out, "in", "ВХОДНЫЕ", sections, log)
        results.append(out)
    if template_out:
        out = os.path.join(out_dir, os.path.basename(template_out))
        log("Сборка перечня выходных сигналов...")
        build_docx(template_out, out, "out", "ВЫХОДНЫЕ", sections, log)
        results.append(out)
    if update_fields and results:
        log("Обновление номеров страниц через Word...")
        if update_fields_word(results, log, pdf=make_pdf):
            log("ГОТОВО.")
            return results
    log("ГОТОВО. При открытии в Word ответьте «Да» на обновление полей")
    log("(или Ctrl+A, F9) — проставятся номера страниц в содержании.")
    return results
